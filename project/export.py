from __future__ import annotations

import json
import re
from dataclasses import dataclass
from datetime import UTC, datetime, timedelta, timezone, tzinfo
from pathlib import Path
from typing import Any, Literal
from zoneinfo import ZoneInfo, ZoneInfoNotFoundError

from project.models import CrawlResult, Video

ExportFormat = Literal["docx", "markdown", "text"]

NAVY = (0x0B, 0x25, 0x45)
BLUE = (0x2E, 0x74, 0xB5)
MUTED = (0x66, 0x6D, 0x75)
DECORATIVE_SYMBOLS = re.compile(
    r"[\U0001F300-\U0001FAFF\U00002600-\U000027BF\uFFFD]"
)
SHANGHAI_TZ: tzinfo
try:
    SHANGHAI_TZ = ZoneInfo("Asia/Shanghai")
except ZoneInfoNotFoundError:
    # Windows installations may not ship the IANA tzdata database.
    SHANGHAI_TZ = timezone(timedelta(hours=8))


class ExportError(RuntimeError):
    """Raised when a result cannot be exported."""


class ExportDependencyError(ExportError):
    """Raised when an optional exporter dependency is not installed."""


@dataclass(frozen=True)
class ExportReport:
    output_path: Path
    run_dir: Path
    format: ExportFormat
    videos: int


@dataclass(frozen=True)
class ExportItem:
    number: int
    video: Video
    raw_text: str
    clean_text: str


def export_result(
    result_path: Path,
    *,
    output_path: Path | None = None,
    output_format: ExportFormat = "docx",
    run_dir: Path | None = None,
) -> ExportReport:
    """Export a saved crawl result and preserve raw/clean transcript sidecars."""

    result_path = result_path.expanduser().resolve()
    if not result_path.is_file():
        raise ExportError(f"结果文件不存在: {result_path}")
    try:
        payload = json.loads(result_path.read_text(encoding="utf-8"))
        result = CrawlResult.model_validate(payload)
    except (OSError, json.JSONDecodeError, ValueError) as exc:
        raise ExportError(f"无法读取结果文件: {result_path}: {exc}") from exc

    output = (
        output_path.expanduser().resolve()
        if output_path is not None
        else default_output_path(result_path, result.user.nickname, output_format)
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    export_run_dir = _make_run_dir(run_dir or output.parent / "rewrites")
    items = _prepare_items(result, result_path, export_run_dir)

    if output_format == "docx":
        _write_docx(output, result, items)
    elif output_format == "markdown":
        _write_markdown(output, result, items)
    elif output_format == "text":
        _write_text(output, result, items)
    else:
        raise ExportError(f"不支持的导出格式: {output_format}")

    return ExportReport(
        output_path=output,
        run_dir=export_run_dir,
        format=output_format,
        videos=len(items),
    )


def default_output_path(result_path: Path, nickname: str, output_format: ExportFormat) -> Path:
    extension = {"docx": "docx", "markdown": "md", "text": "txt"}[output_format]
    safe_name = safe_filename(nickname) or "douyin"
    return result_path.parent / f"{safe_name}_全部口播文案.{extension}"


def clean_transcript(text: str) -> str:
    """Remove timestamps, standalone sound labels, and obvious transcript noise."""

    text = text.replace("\ufeff", "").replace("\r\n", "\n").replace("\r", "\n")
    text = DECORATIVE_SYMBOLS.sub("", text)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    lines: list[str] = []
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if re.fullmatch(
            r"[\[\(]?\s*\d{1,2}:\d{2}(?::\d{2})?(?:[.,]\d{1,3})?\s*[\]\)]?",
            line,
        ):
            continue
        if re.fullmatch(
            r"[\[\(]?\s*(音乐|背景音乐|笑声|掌声|音效|杂音|music|laughter)\s*[\]\)]?",
            line,
            re.IGNORECASE,
        ):
            continue
        line = re.sub(
            r"(?:\[|\()\s*\d{1,2}:\d{2}(?::\d{2})?(?:[.,]\d{1,3})?\s*(?:\]|\))",
            "",
            line,
        )
        line = re.sub(
            r"(?:\[|\()\s*(?:音乐|背景音乐|笑声|掌声|音效|杂音|music|laughter)\s*(?:\]|\))",
            "",
            line,
            flags=re.IGNORECASE,
        )
        line = re.sub(r"^\s*(?:嗯+|呃+|额+)[\s，,、:：]*", "", line)
        line = re.sub(
            r"(^|[，,。！？；;：:])\s*(?:嗯+|呃+|额+)(?=[，,。！？；;：:])",
            r"\1",
            line,
        )
        line = re.sub(r"[ \t]+", " ", line).strip()
        line = re.sub(r"([。！？；])\1+", r"\1", line)
        if line and (not lines or line != lines[-1]):
            lines.append(line)
    cleaned = "\n".join(lines).strip()
    return cleaned or text.strip()


def format_publish_time(value: datetime | None) -> str:
    if value is None:
        return "未提供"
    return value.astimezone(SHANGHAI_TZ).strftime("%Y-%m-%d %H:%M")


def normalize_title(title: str) -> str:
    cleaned = DECORATIVE_SYMBOLS.sub("", title.replace("\ufeff", ""))
    return re.sub(r"\s+", " ", cleaned).strip() or "（无标题）"


def safe_filename(value: str) -> str:
    value = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", value)
    return value.strip(" .")


def _prepare_items(result: CrawlResult, result_path: Path, run_dir: Path) -> list[ExportItem]:
    videos = sorted(
        result.videos,
        key=lambda video: video.publish_time or datetime.min.replace(tzinfo=UTC),
        reverse=True,
    )
    videos_dir = run_dir / "videos"
    items: list[ExportItem] = []
    for number, video in enumerate(videos, start=1):
        raw_text = _read_transcript(video, result_path)
        clean_text = clean_transcript(raw_text)
        item_dir = videos_dir / (safe_filename(video.aweme_id) or f"item-{number:04d}")
        item_dir.mkdir(parents=True, exist_ok=True)
        (item_dir / "transcript_raw.txt").write_text(raw_text, encoding="utf-8")
        (item_dir / "transcript_clean.txt").write_text(clean_text, encoding="utf-8")
        items.append(
            ExportItem(
                number=number,
                video=video,
                raw_text=raw_text,
                clean_text=clean_text,
            )
        )
    return items


def _read_transcript(video: Video, result_path: Path) -> str:
    if video.transcription is None:
        return ""
    transcript_path = Path(video.transcription.transcript_file).expanduser()
    if not transcript_path.is_absolute():
        transcript_path = result_path.parent / transcript_path
    try:
        return transcript_path.read_text(encoding="utf-8") if transcript_path.is_file() else ""
    except OSError as exc:
        raise ExportError(f"无法读取转写文件: {transcript_path}: {exc}") from exc


def _make_run_dir(root: Path) -> Path:
    root = root.expanduser().resolve()
    root.mkdir(parents=True, exist_ok=True)
    stem = datetime.now().strftime("%Y%m%d-%H%M%S")
    candidate = root / stem
    suffix = 1
    while candidate.exists():
        candidate = root / f"{stem}-{suffix:02d}"
        suffix += 1
    candidate.mkdir(parents=True)
    return candidate


def _write_markdown(path: Path, result: CrawlResult, items: list[ExportItem]) -> None:
    lines = [
        f"# {normalize_title(result.user.nickname)}_全部口播文案",
        "",
        f"- 账号：{result.user.nickname}",
        f"- 作品数量：{len(items)}",
        "",
    ]
    for item in items:
        video = item.video
        lines.extend(
            (
                f"## {item.number}. {normalize_title(video.title)}",
                "",
                f"发布时间：{format_publish_time(video.publish_time)}",
                f"视频链接：{video.video_url or '未提供'}",
                "",
                "口播文案：",
                "",
                item.clean_text or "（无转写内容）",
                "",
                "---",
                "",
            )
        )
    path.write_text("\n".join(lines), encoding="utf-8")


def _write_text(path: Path, result: CrawlResult, items: list[ExportItem]) -> None:
    lines = [
        f"{normalize_title(result.user.nickname)}_全部口播文案",
        f"账号：{result.user.nickname}",
        f"作品数量：{len(items)}",
        "",
    ]
    for item in items:
        video = item.video
        lines.extend(
            (
                f"{item.number}. {normalize_title(video.title)}",
                f"发布时间：{format_publish_time(video.publish_time)}",
                f"视频链接：{video.video_url or '未提供'}",
                "",
                "口播文案：",
                item.clean_text or "（无转写内容）",
                "",
                "────────────────",
                "",
            )
        )
    path.write_text("\n".join(lines), encoding="utf-8")


def _write_docx(path: Path, result: CrawlResult, items: list[ExportItem]) -> None:
    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ModuleNotFoundError as exc:
        raise ExportDependencyError(
            "Word 导出需要 python-docx，请安装 douyin-ingest[word]"
        ) from exc

    def set_run_font(
        run: Any,
        name: str = "Calibri",
        east_asia: str = "Microsoft YaHei",
        size: float | None = None,
        color: Any | None = None,
        bold: bool | None = None,
    ) -> None:
        run.font.name = name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), name)
        rfonts.set(qn("w:hAnsi"), name)
        rfonts.set(qn("w:eastAsia"), east_asia)
        if size is not None:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
        if bold is not None:
            run.bold = bold

    def set_style_font(
        style: Any,
        size: float,
        color: Any | None = None,
        bold: bool | None = None,
    ) -> None:
        style.font.name = "Calibri"
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), "Calibri")
        rfonts.set(qn("w:hAnsi"), "Calibri")
        rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        if color is not None:
            style.font.color.rgb = color
        if bold is not None:
            style.font.bold = bold

    def add_hyperlink(paragraph: Any, text: str, url: str) -> None:
        relationship_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "2E74B5")
        rpr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        rpr.append(underline)
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), "Calibri")
        rfonts.set(qn("w:hAnsi"), "Calibri")
        rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        rpr.append(rfonts)
        run.append(rpr)
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(text_node)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def add_page_field(paragraph: Any) -> None:
        run = paragraph.add_run()
        set_run_font(run, size=9, color=RGBColor(*MUTED))
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, instruction, end])

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    if "Video Heading" in [style.name for style in doc.styles]:
        video_heading = doc.styles["Video Heading"]
    else:
        video_heading = doc.styles.add_style("Video Heading", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(video_heading, 13, RGBColor(*NAVY), True)
    video_heading.paragraph_format.space_before = Pt(14)
    video_heading.paragraph_format.space_after = Pt(5)
    video_heading.paragraph_format.line_spacing = 1.15
    video_heading.paragraph_format.keep_with_next = True

    if "Transcript Body" in [style.name for style in doc.styles]:
        transcript_body = doc.styles["Transcript Body"]
    else:
        transcript_body = doc.styles.add_style("Transcript Body", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(transcript_body, 10.5)
    transcript_body.paragraph_format.space_before = Pt(0)
    transcript_body.paragraph_format.space_after = Pt(6)
    transcript_body.paragraph_format.line_spacing = 1.25
    transcript_body.paragraph_format.widow_control = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    footer.clear()
    footer_label = footer.add_run(f"{result.user.nickname} · 第 ")
    set_run_font(footer_label, size=9, color=RGBColor(*MUTED))
    add_page_field(footer)
    footer_end = footer.add_run(" 页")
    set_run_font(footer_end, size=9, color=RGBColor(*MUTED))

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(7)
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run(f"{normalize_title(result.user.nickname)}_全部口播文案")
    set_run_font(title_run, size=23, color=RGBColor(*NAVY), bold=True)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    meta_run = meta.add_run(f"账号：{result.user.nickname}  ·  全部作品：{len(items)} 条")
    set_run_font(meta_run, size=10, color=RGBColor(*MUTED))

    for item in items:
        video = item.video
        heading = doc.add_paragraph(style="Video Heading")
        heading.paragraph_format.keep_with_next = True
        heading_run = heading.add_run(f"{item.number}. {normalize_title(video.title)}")
        set_run_font(heading_run, size=13, color=RGBColor(*NAVY), bold=True)

        _add_docx_label(
            doc,
            "发布时间：",
            format_publish_time(video.publish_time),
            set_run_font,
            Pt,
        )
        video_url = video.video_url or ""
        link_paragraph = doc.add_paragraph()
        link_paragraph.paragraph_format.space_after = Pt(3)
        label_run = link_paragraph.add_run("视频链接：")
        set_run_font(label_run, size=10.5, color=RGBColor(*NAVY), bold=True)
        if video_url:
            add_hyperlink(link_paragraph, video_url, video_url)
        else:
            value_run = link_paragraph.add_run("未提供")
            set_run_font(value_run, size=10.5)

        speech_label = doc.add_paragraph()
        speech_label.paragraph_format.space_before = Pt(5)
        speech_label.paragraph_format.space_after = Pt(3)
        speech_label.paragraph_format.keep_with_next = True
        speech_run = speech_label.add_run("口播文案：")
        set_run_font(speech_run, size=10.5, color=RGBColor(*NAVY), bold=True)

        blocks = item.clean_text.split("\n") if item.clean_text else ["（无转写内容）"]
        for block in blocks:
            body = doc.add_paragraph(style="Transcript Body")
            body_run = body.add_run(block)
            set_run_font(body_run, size=10.5)

        if item.number < len(items):
            separator = doc.add_paragraph()
            separator.paragraph_format.space_before = Pt(7)
            separator.paragraph_format.space_after = Pt(2)
            separator.paragraph_format.keep_together = True
            separator_run = separator.add_run("────────────────")
            set_run_font(separator_run, size=9, color=RGBColor(*MUTED))

    doc.core_properties.title = f"{result.user.nickname}_全部口播文案"
    doc.core_properties.subject = "抖音主页视频中文口播转写汇编"
    doc.save(str(path))


def _add_docx_label(
    doc: Any,
    label: str,
    value: str,
    set_run_font: Any,
    pt: Any,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = pt(0)
    paragraph.paragraph_format.space_after = pt(3)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, size=10.5, color=_docx_color(0x0B, 0x25, 0x45), bold=True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, size=10.5)


def _docx_color(red: int, green: int, blue: int) -> Any:
    from docx.shared import RGBColor

    return RGBColor(red, green, blue)
