#!/usr/bin/env python3
"""Build and structurally verify the fixed Douyin rewrite Word deliverable."""

from __future__ import annotations

import argparse
import html
import json
import re
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any
from zipfile import ZipFile

TEMPLATE_VERSION = "douyin-script-rewriter-word-v1"
FONT_NAME = "Arial Unicode MS"
ACCENT_RED = "C9362B"
ACCENT_GREEN = "287D5A"


try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ModuleNotFoundError as error:
    DOCX_IMPORT_ERROR: ModuleNotFoundError | None = error
else:
    DOCX_IMPORT_ERROR = None


def require_docx_dependency() -> None:
    if DOCX_IMPORT_ERROR is not None:
        raise RuntimeError(
            "缺少 python-docx。请先运行：python -m pip install 'douyin-ingest[word]'；"
            "源码开发环境可运行：python -m pip install -e '.[word]'。"
        ) from DOCX_IMPORT_ERROR


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="从 douyin-script-rewriter 的 result.json 生成固定模板 Word。"
    )
    parser.add_argument(
        "--run",
        required=True,
        help="rewrite 运行目录，或该目录内的 result.json/report.md。",
    )
    parser.add_argument(
        "--output",
        help="可选 DOCX 输出路径；默认写入 <run>/deliverables/。",
    )
    parser.add_argument(
        "--verify-only",
        action="store_true",
        help="不重新生成，只校验 --output 或默认 DOCX。",
    )
    parser.add_argument(
        "--no-record",
        action="store_true",
        help="不把 word_file 写回 result.json，也不更新 report.md。用于回归测试。",
    )
    return parser.parse_args()


def resolve_run(value: str) -> Path:
    path = Path(value).expanduser().resolve()
    if path.is_file():
        if path.name not in {"result.json", "report.md"}:
            raise ValueError("--run 文件必须是 result.json 或 report.md")
        path = path.parent
    if not path.is_dir():
        raise FileNotFoundError(f"运行目录不存在：{path}")
    if not (path / "result.json").is_file():
        raise FileNotFoundError(f"缺少 result.json：{path / 'result.json'}")
    return path


def load_result(run: Path) -> dict[str, Any]:
    result = json.loads((run / "result.json").read_text(encoding="utf-8"))
    source = result.get("source")
    videos = result.get("videos")
    if not isinstance(source, dict) or not isinstance(videos, list) or not videos:
        raise ValueError("result.json 缺少有效的 source 或 videos[]")

    for index, video in enumerate(videos, start=1):
        if not isinstance(video, dict):
            raise ValueError(f"videos[{index - 1}] 不是对象")
        if not video.get("title") and video.get("name"):
            video["title"] = video["name"]
        if not video.get("source_url"):
            video["source_url"] = video.get("video_url") or video.get("page_url")
        if not video.get("status"):
            video["status"] = "failed" if video.get("error") else "success"
        for field in ("title", "source_url", "digg_count", "status"):
            if video.get(field) in (None, ""):
                raise ValueError(f"第 {index} 条缺少字段：{field}")
        rank = video.get("rank", index)
        if rank != index:
            raise ValueError(f"视频顺序与 rank 不一致：第 {index} 条的 rank={rank}")
        video["rank"] = rank

        if video["status"] == "success":
            clean_path = video.get("clean_transcript_file")
            if not clean_path or not Path(clean_path).is_file():
                raise FileNotFoundError(f"TOP {index:02d} 缺少 AI 校正版逐字稿")
            rewrite_path = video.get("rewrite_file")
            if rewrite_path and not Path(rewrite_path).is_file():
                raise FileNotFoundError(f"TOP {index:02d} 的 rewrite_file 不存在")
        elif not video.get("error"):
            raise ValueError(f"TOP {index:02d} 处理失败但没有 error")
    if not result.get("status"):
        success_count = sum(video["status"] == "success" for video in videos)
        if success_count == len(videos):
            result["status"] = "success"
        elif success_count:
            result["status"] = "partial"
        else:
            result["status"] = "failed"
    return result


def safe_filename(value: str) -> str:
    value = re.sub(r"[\\/:*?\"<>|\x00-\x1f]", "-", value).strip(" .")
    value = re.sub(r"\s+", "", value)
    return value[:80] or "抖音内容"


def scope_label(result: dict[str, Any]) -> str:
    videos = result["videos"]
    if result["source"].get("collection_mode") == "profile" or len(videos) > 1:
        return f"Top{len(videos)}"
    return "单条"


def scope_display(result: dict[str, Any]) -> str:
    videos = result["videos"]
    if result["source"].get("collection_mode") == "profile" or len(videos) > 1:
        return f"Top {len(videos)}"
    return "单条"


def account_name(result: dict[str, Any]) -> str:
    name = result["source"].get("account_name")
    if name:
        return str(name)
    return str(result["videos"][0]["title"])


def default_output_path(run: Path, result: dict[str, Any]) -> Path:
    filename = (
        f"{safe_filename(account_name(result))}-{scope_label(result)}内容改写-简洁版.docx"
    )
    return run / "deliverables" / filename


def read_artifact(path_value: str | None) -> str | None:
    if not path_value:
        return None
    path = Path(path_value)
    text = path.read_text(encoding="utf-8").strip()
    if not text:
        raise ValueError(f"内容文件为空：{path}")

    if path.suffix.lower() == ".md" and text.startswith("# "):
        text = text.split("\n", 1)[1].lstrip() if "\n" in text else ""

    heading_lines = {
        "AI 校正版逐字稿",
        "原创口播稿",
        "原创改写稿",
        "改写稿",
    }
    lines = text.splitlines()
    while lines and lines[0].strip().lstrip("#").strip() in heading_lines:
        lines.pop(0)
        while lines and not lines[0].strip():
            lines.pop(0)
    text = "\n".join(lines).strip()
    if not text:
        raise ValueError(f"去除重复标题后内容为空：{path}")
    return text


def derive_date(run: Path) -> str:
    match = re.match(r"(\d{4})(\d{2})(\d{2})", run.name)
    if match:
        return "-".join(match.groups())
    return datetime.now().astimezone().strftime("%Y-%m-%d")


def set_font(
    run: Any,
    size: float,
    color: str = "1C1C1C",
    *,
    bold: bool = False,
) -> None:
    run.font.name = FONT_NAME
    fonts = run._element.get_or_add_rPr().rFonts
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def set_style_font(style: Any, size: float, color: str, *, bold: bool) -> None:
    style.font.name = FONT_NAME
    fonts = style._element.get_or_add_rPr().rFonts
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), FONT_NAME)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT_RED)
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    fonts = OxmlElement("w:rFonts")
    for name in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{name}"), FONT_NAME)
    properties.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    properties.append(size)

    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, 9, "777777")
    for tag, attrs, value in (
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
        ("w:fldChar", {"w:fldCharType": "separate"}, None),
        ("w:t", {}, "1"),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ):
        node = OxmlElement(tag)
        for key, attr_value in attrs.items():
            node.set(qn(key), attr_value)
        if value is not None:
            node.text = value
        run._r.append(node)
    set_font(paragraph.add_run(" 页"), 9, "777777")


def add_script_paragraphs(document: Document, text: str) -> None:
    blocks = [part.strip() for part in text.split("\n\n") if part.strip()]
    for block in blocks:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.keep_together = False
        paragraph.paragraph_format.widow_control = True
        for index, line in enumerate(block.splitlines()):
            if index:
                paragraph.add_run().add_break()
            set_font(paragraph.add_run(line), 11)


def configure_document(document: Document, name: str, scope: str) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    set_style_font(normal, 11, "1C1C1C", bold=False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_1 = document.styles["Heading 1"]
    set_style_font(heading_1, 16, ACCENT_RED, bold=True)
    heading_1.paragraph_format.space_before = Pt(18)
    heading_1.paragraph_format.space_after = Pt(10)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = document.styles["Heading 2"]
    set_style_font(heading_2, 13, ACCENT_GREEN, bold=True)
    heading_2.paragraph_format.space_before = Pt(14)
    heading_2.paragraph_format.space_after = Pt(7)
    heading_2.paragraph_format.keep_with_next = True

    metadata = document.styles.add_style("Metadata", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(metadata, 9.5, "666666", bold=False)
    metadata.paragraph_format.space_after = Pt(4)
    metadata.paragraph_format.line_spacing = 1.15

    warning = document.styles.add_style("Warning", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(warning, 9.5, "5C4A23", bold=False)
    warning.paragraph_format.left_indent = Inches(0.18)
    warning.paragraph_format.space_before = Pt(4)
    warning.paragraph_format.space_after = Pt(8)
    warning.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run(f"{name}  |  {scope} 内容改写"), 8.5, "777777", bold=True)
    add_page_number(section.footer.paragraphs[0])


def build_document(run: Path, result: dict[str, Any], output: Path) -> None:
    require_docx_dependency()
    name = account_name(result)
    scope = scope_display(result)
    videos = result["videos"]
    document = Document()
    configure_document(document, name, scope)

    for _ in range(4):
        document.add_paragraph()
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_font(kicker.add_run("DOUYIN CONTENT BOOK"), 10, ACCENT_RED, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run(name), 28, "202326", bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    set_font(subtitle.add_run(f"点赞 {scope} 内容改写"), 16, "3F454A")

    successful = [video for video in videos if video["status"] == "success"]
    has_rewrite = any(video.get("rewrite_file") for video in successful)
    content_label = "AI 校正版逐字稿 + 原创口播稿" if has_rewrite else "AI 校正版逐字稿"
    description = document.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.paragraph_format.space_after = Pt(50)
    set_font(description.add_run(content_label), 11, "777777")

    date_paragraph = document.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(date_paragraph.add_run(derive_date(run)), 10, "777777", bold=True)

    if len(videos) > 1:
        document.add_page_break()
        document.add_heading("内容目录", level=1)
        intro = document.add_paragraph()
        set_font(
            intro.add_run(
                f"按点赞量从高到低排列，共 {len(videos)} 条。点击标题下方链接可打开抖音原视频。"
            ),
            10.5,
            "555555",
        )
        for video in videos:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(5)
            set_font(
                paragraph.add_run(f"{video['rank']:02d}  "),
                10.5,
                ACCENT_RED,
                bold=True,
            )
            set_font(paragraph.add_run(str(video["title"])), 10.5, "1C1C1C", bold=True)
            set_font(
                paragraph.add_run(f"  ·  {int(video['digg_count']):,} 赞"),
                9.5,
                "777777",
            )

    for video in videos:
        document.add_page_break()
        rank = int(video["rank"])
        rank_paragraph = document.add_paragraph()
        rank_paragraph.paragraph_format.space_after = Pt(4)
        set_font(rank_paragraph.add_run(f"TOP {rank:02d}"), 10, ACCENT_RED, bold=True)

        heading = document.add_heading(str(video["title"]), level=1)
        heading.paragraph_format.space_before = Pt(0)

        metadata = document.add_paragraph(style="Metadata")
        set_font(
            metadata.add_run(f"点赞量：{int(video['digg_count']):,}    "),
            9.5,
            "666666",
            bold=True,
        )
        add_hyperlink(metadata, "打开抖音原视频", str(video["source_url"]))

        if video["status"] != "success":
            document.add_heading("处理状态", level=2)
            paragraph = document.add_paragraph(style="Warning")
            set_font(paragraph.add_run(f"处理失败：{video['error']}"), 9.5, "5C4A23")
            continue

        if video.get("quality_grade") == "unreliable":
            paragraph = document.add_paragraph(style="Warning")
            set_font(
                paragraph.add_run(
                    "转写提示：基础模型无法确认的片段已标为〔听不清〕或〔疑似〕，"
                    "改写仅采用可确认含义。"
                ),
                9.5,
                "5C4A23",
            )

        clean_text = read_artifact(video.get("clean_transcript_file"))
        if clean_text is None:
            raise ValueError(f"TOP {rank:02d} 缺少 AI 校正版逐字稿")
        document.add_heading("AI 校正版逐字稿", level=2)
        add_script_paragraphs(document, clean_text)

        rewrite_text = read_artifact(video.get("rewrite_file"))
        if rewrite_text is not None:
            document.add_heading("原创口播稿", level=2)
            add_script_paragraphs(document, rewrite_text)

    properties = document.core_properties
    properties.title = f"{name}-{scope_label(result)}内容改写"
    properties.subject = "抖音热门内容 AI 校正与原创改写"
    properties.identifier = TEMPLATE_VERSION
    properties.keywords = TEMPLATE_VERSION
    properties.author = ""
    properties.last_modified_by = ""
    properties.comments = ""

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def compact_text(value: str) -> str:
    return re.sub(r"\s+", "", html.unescape(value))


def verify_document(result: dict[str, Any], output: Path) -> dict[str, Any]:
    if not output.is_file() or output.stat().st_size < 10_000:
        raise ValueError(f"DOCX 不存在或文件过小：{output}")

    with ZipFile(output) as archive:
        required_members = {
            "[Content_Types].xml",
            "word/document.xml",
            "word/_rels/document.xml.rels",
            "docProps/core.xml",
        }
        missing = required_members - set(archive.namelist())
        if missing:
            raise ValueError(f"DOCX 结构不完整：{sorted(missing)}")
        document_xml = archive.read("word/document.xml").decode("utf-8")
        relationships_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        core_xml = archive.read("docProps/core.xml").decode("utf-8")

    paragraph_xml = re.findall(r"<w:p(?:\s[^>]*)?>.*?</w:p>", document_xml, re.DOTALL)
    paragraph_texts = [
        html.unescape(
            "".join(re.findall(r"<w:t(?:\s[^>]*)?>(.*?)</w:t>", paragraph))
        )
        for paragraph in paragraph_xml
    ]
    text_nodes = re.findall(r"<w:t(?:\s[^>]*)?>(.*?)</w:t>", document_xml)
    document_text = html.unescape("".join(text_nodes))
    compact_document = compact_text(document_text)
    videos = result["videos"]

    if "{{" in document_text or "}}" in document_text:
        raise ValueError("DOCX 中残留模板占位符")
    if TEMPLATE_VERSION not in core_xml:
        raise ValueError("DOCX 模板版本标记缺失")

    name = account_name(result)
    scope = scope_display(result)
    successful = [video for video in videos if video["status"] == "success"]
    has_rewrite = any(video.get("rewrite_file") for video in successful)
    content_label = "AI 校正版逐字稿 + 原创口播稿" if has_rewrite else "AI 校正版逐字稿"
    cover_requirements = ("DOUYIN CONTENT BOOK", name, f"点赞 {scope} 内容改写", content_label)
    for required_text in cover_requirements:
        if compact_text(required_text) not in compact_document:
            raise ValueError(f"封面固定内容缺失：{required_text}")

    expected_directory_count = 1 if len(videos) > 1 else 0
    if paragraph_texts.count("内容目录") != expected_directory_count:
        raise ValueError("内容目录数量与交付范围不匹配")
    if len(videos) > 1:
        directory_intro = (
            f"按点赞量从高到低排列，共 {len(videos)} 条。"
            "点击标题下方链接可打开抖音原视频。"
        )
        if compact_text(directory_intro) not in compact_document:
            raise ValueError("内容目录说明缺失或发生变化")

    labels = [
        paragraph
        for paragraph in paragraph_texts
        if re.fullmatch(r"TOP (?:0[1-9]|[1-9][0-9]+)", paragraph)
    ]
    expected_labels = [f"TOP {int(video['rank']):02d}" for video in videos]
    if labels != expected_labels:
        raise ValueError(f"TOP 标签不匹配：expected={expected_labels}, actual={labels}")

    expected_clean = 0
    expected_rewrite = 0
    for video in videos:
        if compact_text(str(video["title"])) not in compact_document:
            raise ValueError(f"缺少视频标题：{video['title']}")
        likes = int(video["digg_count"])
        if compact_text(f"点赞量：{likes:,}") not in compact_document:
            raise ValueError(f"缺少点赞量：TOP {int(video['rank']):02d}")
        if len(videos) > 1 and compact_text(f"{likes:,} 赞") not in compact_document:
            raise ValueError(f"目录缺少点赞量：TOP {int(video['rank']):02d}")
        if video["status"] != "success":
            failure_text = f"处理失败：{video['error']}"
            if compact_text(failure_text) not in compact_document:
                raise ValueError(f"失败状态内容不完整：TOP {int(video['rank']):02d}")
            continue
        clean_text = read_artifact(video.get("clean_transcript_file"))
        if clean_text is None or compact_text(clean_text) not in compact_document:
            raise ValueError(f"AI 校正版内容不完整：TOP {int(video['rank']):02d}")
        expected_clean += 1
        rewrite_text = read_artifact(video.get("rewrite_file"))
        if rewrite_text is not None:
            if compact_text(rewrite_text) not in compact_document:
                raise ValueError(f"原创口播稿内容不完整：TOP {int(video['rank']):02d}")
            expected_rewrite += 1

    if paragraph_texts.count("AI 校正版逐字稿") != expected_clean:
        raise ValueError("AI 校正版逐字稿标题数量不匹配或发生重复")
    if paragraph_texts.count("原创口播稿") != expected_rewrite:
        raise ValueError("原创口播稿标题数量不匹配或发生重复")

    actual_links = [
        html.unescape(value)
        for value in re.findall(r'Target="(https?://[^\"]+)"', relationships_xml)
    ]
    expected_links = [str(video["source_url"]) for video in videos]
    if Counter(actual_links) != Counter(expected_links):
        raise ValueError(
            f"原视频链接不匹配：expected={len(expected_links)}, actual={len(actual_links)}"
        )

    return {
        "ok": True,
        "template_version": TEMPLATE_VERSION,
        "word_file": str(output),
        "video_count": len(videos),
        "source_links": len(actual_links),
        "clean_sections": expected_clean,
        "rewrite_sections": expected_rewrite,
        "bytes": output.stat().st_size,
    }


def record_deliverable(run: Path, result: dict[str, Any], output: Path) -> None:
    result["schema_version"] = "1.1"
    result["word_file"] = str(output.resolve())
    result["word_template_version"] = TEMPLATE_VERSION
    result_path = run / "result.json"
    temporary = result_path.with_suffix(".json.tmp")
    temporary.write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    temporary.replace(result_path)

    report_path = run / "report.md"
    if not report_path.is_file():
        return
    report = report_path.read_text(encoding="utf-8")
    link = output.relative_to(run).as_posix() if output.is_relative_to(run) else str(output)
    line = f"**Word 版：[{output.name}]({link})**"
    pattern = re.compile(r"^\*\*Word 版：.*?\*\*$", re.MULTILINE)
    if pattern.search(report):
        report = pattern.sub(line, report, count=1)
    else:
        lines = report.splitlines()
        insert_at = 1 if lines and lines[0].startswith("# ") else 0
        lines[insert_at:insert_at] = ["", line]
        report = "\n".join(lines)
    report_path.write_text(report.rstrip() + "\n", encoding="utf-8")


def main() -> int:
    args = parse_args()
    run = resolve_run(args.run)
    result = load_result(run)
    output = (
        Path(args.output).expanduser().resolve()
        if args.output
        else default_output_path(run, result).resolve()
    )

    if not args.verify_only:
        build_document(run, result, output)
    summary = verify_document(result, output)
    if not args.verify_only and not args.no_record:
        record_deliverable(run, result, output)
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except (FileNotFoundError, RuntimeError, ValueError, json.JSONDecodeError) as error:
        print(json.dumps({"ok": False, "error": str(error)}, ensure_ascii=False), file=sys.stderr)
        raise SystemExit(1) from error
